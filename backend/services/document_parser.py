from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class DocumentParser:

    def extract_docx_text(self, file_path: Path) -> str:
        """Extracts transient plain text from uploaded DOCX files (python-docx)."""
        try:
            document = Document(str(file_path))
        except PackageNotFoundError as exc:
            # python-docx rejects anything that isn't a valid OOXML zip (e.g. a
            # legacy .doc, an encrypted doc, or a mislabeled/corrupt file).
            raise ValueError("File is not a valid .docx document") from exc
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(part for part in parts if part.strip())
